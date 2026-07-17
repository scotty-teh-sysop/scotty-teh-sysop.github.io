#!/usr/bin/env python3
"""
Alberta Court of King's Bench (CKB) Document Sorter & Formatter
================================================================

This tool parses, validates, and sorts legal documents according to
Alberta Court of King's Bench formatting rules for all jurisdictions in Alberta.

Key CKB Header Fields (per Alberta Rules of Court):
- COURT FILE NUMBER (e.g., 2303-12345)
- JUDICIAL CENTRE (Edmonton, Calgary, Lethbridge, Medicine Hat, Red Deer, St. Paul, Wetaskiwin)
- PLAINTIFF / DEFENDANT (civil) or PROSECUTION / ACCUSED (criminal)
- DOCUMENT TYPE (Statement of Claim, Affidavit, Order, Application, etc.)
- ADDRESS FOR SERVICE AND CONTACT INFORMATION

Sort Order:
1. By JUDICIAL CENTRE (alphabetical: Calgary, Edmonton, Lethbridge, Medicine Hat, Red Deer, St. Paul, Wetaskiwin)
2. By DATE ID (COURT FILE NUMBER year/sequence)
3. By CASE ID (sequential)
4. By PARTY TYPE (Prosecution/Plaintiff first, then Defendant/Accused)
5. By DOCUMENT TYPE (per CKB precedence)
6. By JURISDICTION (Civil, Criminal, Family, Surrogate, Bankruptcy)

Usage:
    python ckb_document_sorter.py [input_dir] [output_dir] [--format md|json|html]
"""

import os
import re
import json
import shutil
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional, List, Dict, Any
from enum import Enum
import argparse


class Jurisdiction(Enum):
    CIVIL = "civil"
    CRIMINAL = "criminal"
    FAMILY = "family"
    SURROGATE = "surrogate"
    BANKRUPTCY = "bankruptcy"
    UNKNOWN = "unknown"


class PartyType(Enum):
    PROSECUTION = "prosecution"
    PLAINTIFF = "plaintiff"
    DEFENDANT = "defendant"
    ACCUSED = "accused"
    APPLICANT = "applicant"
    RESPONDENT = "respondent"
    UNKNOWN = "unknown"


class DocumentType(Enum):
    # Civil precedence (per Alberta Rules of Court)
    ORIGINATING_APPLICATION = "originating_application"
    STATEMENT_OF_CLAIM = "statement_of_claim"
    STATEMENT_OF_DEFENCE = "statement_of_defence"
    COUNTERCLAIM = "counterclaim"
    REPLY = "reply"
    AFFIDAVIT = "affidavit"
    APPLICATION = "application"
    ORDER = "order"
    JUDGMENT = "judgment"
    NOTICE_OF_APPEAL = "notice_of_appeal"
    # Criminal
    INFORMATION = "information"
    INDICTMENT = "indictment"
    BAIL_ORDER = "bail_order"
    SENTENCING_ORDER = "sentencing_order"
    # Family
    CLAIM = "claim"
    RESPONSE = "response"
    PARENTING_ORDER = "parenting_order"
    SUPPORT_ORDER = "support_order"
    # General
    NOTICE = "notice"
    MOTION = "motion"
    BRIEF = "brief"
    TRANSCRIPT = "transcript"
    EXHIBIT = "exhibit"
    CORRESPONDENCE = "correspondence"
    OTHER = "other"


# Judicial Centres in Alberta (alphabetical sort order)
JUDICIAL_CENTRES = [
    "Calgary",
    "Edmonton",
    "Lethbridge",
    "Medicine Hat",
    "Red Deer",
    "St. Paul",
    "Wetaskiwin",
]

# Document type sort precedence (lower = higher priority)
DOCUMENT_TYPE_PRECEDENCE = {
    DocumentType.ORIGINATING_APPLICATION: 1,
    DocumentType.STATEMENT_OF_CLAIM: 2,
    DocumentType.STATEMENT_OF_DEFENCE: 3,
    DocumentType.COUNTERCLAIM: 4,
    DocumentType.REPLY: 5,
    DocumentType.AFFIDAVIT: 6,
    DocumentType.APPLICATION: 7,
    DocumentType.ORDER: 8,
    DocumentType.JUDGMENT: 9,
    DocumentType.NOTICE_OF_APPEAL: 10,
    DocumentType.INFORMATION: 11,
    DocumentType.INDICTMENT: 12,
    DocumentType.BAIL_ORDER: 13,
    DocumentType.SENTENCING_ORDER: 14,
    DocumentType.CLAIM: 15,
    DocumentType.RESPONSE: 16,
    DocumentType.PARENTING_ORDER: 17,
    DocumentType.SUPPORT_ORDER: 18,
    DocumentType.NOTICE: 19,
    DocumentType.MOTION: 20,
    DocumentType.BRIEF: 21,
    DocumentType.TRANSCRIPT: 22,
    DocumentType.EXHIBIT: 23,
    DocumentType.CORRESPONDENCE: 24,
    DocumentType.OTHER: 99,
}

# Party type sort precedence
PARTY_TYPE_PRECEDENCE = {
    PartyType.PROSECUTION: 1,
    PartyType.PLAINTIFF: 2,
    PartyType.APPLICANT: 3,
    PartyType.DEFENDANT: 4,
    PartyType.ACCUSED: 5,
    PartyType.RESPONDENT: 6,
    PartyType.UNKNOWN: 99,
}

# Jurisdiction sort precedence
JURISDICTION_PRECEDENCE = {
    Jurisdiction.CRIMINAL: 1,
    Jurisdiction.CIVIL: 2,
    Jurisdiction.FAMILY: 3,
    Jurisdiction.SURROGATE: 4,
    Jurisdiction.BANKRUPTCY: 5,
    Jurisdiction.UNKNOWN: 99,
}


@dataclass
class CKBHeader:
    """Parsed CKB document header fields"""
    court_file_number: str = ""
    judicial_centre: str = ""
    plaintiff: str = ""
    defendant: str = ""
    prosecution: str = ""
    accused: str = ""
    document_type: str = ""
    address_for_service: str = ""
    date_filed: str = ""
    clerk_stamp: str = ""
    
    # Computed fields
    jurisdiction: Jurisdiction = Jurisdiction.UNKNOWN
    party_type: PartyType = PartyType.UNKNOWN
    doc_type_enum: DocumentType = DocumentType.OTHER
    date_id: str = ""  # YYYY-MM-DD from file number or date_filed
    case_id: str = ""  # Sequential part of court file number
    year: int = 0
    sequence: int = 0
    
    # Raw header text for reference
    raw_header: str = ""
    
    def __post_init__(self):
        self._parse_computed_fields()
    
    def _parse_computed_fields(self):
        """Extract computed fields from parsed data"""
        # Parse court file number first
        self._parse_court_file_number()
        
        # Determine jurisdiction from document type and parties
        self.jurisdiction = self._determine_jurisdiction()
        
        # Determine party type from who is listed
        self.party_type = self._determine_party_type()
        
        # Map document type string to enum
        self.doc_type_enum = self._map_document_type()
    
    def _parse_court_file_number(self):
        """Parse court file number to extract year and sequence"""
        if not self.court_file_number:
            return
        
        # Alberta Court of King's Bench file number formats:
        # Format 1: YYMM-NNNNN (e.g., 2303-12345 = year 2023, month 03, seq 12345)
        # Format 2: YYYY-NNNNN (e.g., 2023-12345 = year 2023, seq 12345)
        # Format 3: YY-NNNNN (e.g., 23-12345 = year 2023, seq 12345)
        
        clean = self.court_file_number.replace(' ', '').replace('/', '-').replace(':', '-')
        
        # Try YYMM-NNNNN format (6 digits before separator)
        match_yymm = re.search(r'(\d{2})(\d{2})[-:](\d+)', clean)
        if match_yymm:
            year_2digit = match_yymm.group(1)
            month = match_yymm.group(2)
            self.sequence = int(match_yymm.group(3))
            self.case_id = match_yymm.group(3)
            self.year = 2000 + int(year_2digit)
            self.date_id = f"{self.year:04d}-{self.sequence:05d}"
            return
        
        # Try YYYY-NNNNN format (4-digit year)
        match_4yr = re.search(r'(\d{4})[-:](\d+)', clean)
        if match_4yr:
            year_part = match_4yr.group(1)
            self.sequence = int(match_4yr.group(2))
            self.case_id = match_4yr.group(2)
            self.year = int(year_part)
            self.date_id = f"{self.year:04d}-{self.sequence:05d}"
            return
        
        # Try YY-NNNNN format (2-digit year)
        match_2yr = re.search(r'(\d{2})[-:](\d+)', clean)
        if match_2yr:
            year_part = match_2yr.group(1)
            self.sequence = int(match_2yr.group(2))
            self.case_id = match_2yr.group(2)
            self.year = 2000 + int(year_part)
            self.date_id = f"{self.year:04d}-{self.sequence:05d}"
            return
    
    def _determine_jurisdiction(self) -> Jurisdiction:
        doc_lower = self.document_type.lower()
        if any(kw in doc_lower for kw in ['information', 'indictment', 'bail', 'sentencing', 'criminal', 'accused', 'r. v.', 'rex v']):
            return Jurisdiction.CRIMINAL
        if any(kw in doc_lower for kw in ['divorce', 'parenting', 'support', 'family', 'matrimonial', 'child']):
            return Jurisdiction.FAMILY
        if any(kw in doc_lower for kw in ['bankruptcy', 'insolvency', 'receivership']):
            return Jurisdiction.BANKRUPTCY
        if any(kw in doc_lower for kw in ['probate', 'will', 'estate', 'surrogate', 'letters probate']):
            return Jurisdiction.SURROGATE
        if self.prosecution or self.accused:
            return Jurisdiction.CRIMINAL
        return Jurisdiction.CIVIL
    
    def _determine_party_type(self) -> PartyType:
        if self.prosecution:
            return PartyType.PROSECUTION
        if self.plaintiff:
            return PartyType.PLAINTIFF
        if self.defendant:
            return PartyType.DEFENDANT
        if self.accused:
            return PartyType.ACCUSED
        return PartyType.UNKNOWN
    
    def _map_document_type(self) -> DocumentType:
        doc_lower = self.document_type.lower().strip()
        
        # Most specific checks first (avoid substring false positives)
        
        # Family (check before generic "order")
        if 'parenting order' in doc_lower:
            return DocumentType.PARENTING_ORDER
        if 'support order' in doc_lower:
            return DocumentType.SUPPORT_ORDER
        if 'application for grant of probate' in doc_lower or 'grant of probate' in doc_lower:
            return DocumentType.APPLICATION
        if doc_lower == 'claim' and self.jurisdiction == Jurisdiction.FAMILY:
            return DocumentType.CLAIM
        if doc_lower == 'response' and self.jurisdiction == Jurisdiction.FAMILY:
            return DocumentType.RESPONSE
        
        # Civil
        if 'originating application' in doc_lower:
            return DocumentType.ORIGINATING_APPLICATION
        if 'statement of claim' in doc_lower:
            return DocumentType.STATEMENT_OF_CLAIM
        if 'statement of defence' in doc_lower or 'statement of defense' in doc_lower:
            return DocumentType.STATEMENT_OF_DEFENCE
        if 'counterclaim' in doc_lower:
            return DocumentType.COUNTERCLAIM
        if doc_lower == 'reply' or 'reply to' in doc_lower:
            return DocumentType.REPLY
        if 'affidavit' in doc_lower:
            return DocumentType.AFFIDAVIT
        if doc_lower == 'application' or 'notice of application' in doc_lower:
            return DocumentType.APPLICATION
        if 'consent judgment' in doc_lower or 'judgment' in doc_lower:
            return DocumentType.JUDGMENT
        if 'order' in doc_lower and 'consent' not in doc_lower:
            return DocumentType.ORDER
        if 'notice of appeal' in doc_lower:
            return DocumentType.NOTICE_OF_APPEAL
        
        # Criminal
        if 'information' in doc_lower:
            return DocumentType.INFORMATION
        if 'indictment' in doc_lower:
            return DocumentType.INDICTMENT
        if 'bail' in doc_lower:
            return DocumentType.BAIL_ORDER
        if 'sentencing' in doc_lower:
            return DocumentType.SENTENCING_ORDER
        
        # General
        if 'notice' in doc_lower:
            return DocumentType.NOTICE
        if 'motion' in doc_lower:
            return DocumentType.MOTION
        if 'brief' in doc_lower or 'factum' in doc_lower:
            return DocumentType.BRIEF
        if 'transcript' in doc_lower:
            return DocumentType.TRANSCRIPT
        if 'exhibit' in doc_lower:
            return DocumentType.EXHIBIT
        if 'correspondence' in doc_lower or 'letter' in doc_lower:
            return DocumentType.CORRESPONDENCE
        
        return DocumentType.OTHER


@dataclass
class CKBDocument:
    """A complete CKB document with header and content"""
    header: CKBHeader
    content: str = ""
    filepath: str = ""
    filename: str = ""
    
    def sort_key(self) -> tuple:
        """Generate sort key per CKB rules"""
        # 1. Judicial Centre (alphabetical)
        h = self.header
        jc_index = JUDICIAL_CENTRES.index(h.judicial_centre) if h.judicial_centre in JUDICIAL_CENTRES else 99
        
        # 2. Date ID (year-sequence)
        date_id = h.date_id or "9999-99999"
        
        # 3. Case ID (sequence)
        case_id = f"{h.sequence:05d}" if h.sequence else "99999"
        
        # 4. Party Type precedence
        party_prec = PARTY_TYPE_PRECEDENCE.get(h.party_type, 99)
        
        # 5. Document Type precedence
        doc_prec = DOCUMENT_TYPE_PRECEDENCE.get(h.doc_type_enum, 99)
        
        # 6. Jurisdiction precedence
        jur_prec = JURISDICTION_PRECEDENCE.get(h.jurisdiction, 99)
        
        return (jc_index, date_id, case_id, party_prec, doc_prec, jur_prec, self.filename)


class CKBDocumentParser:
    """Parse legal documents to extract CKB header fields"""
    
    # Header field patterns (case-insensitive)
    FIELD_PATTERNS = {
        'court_file_number': [
            r'COURT\s+FILE\s+NUMBER[:\s]*([^\n\r]+)',
            r'COURT\s+FILE\s+NO[.:]\s*([^\n\r]+)',
            r'FILE\s+NUMBER[:\s]*([^\n\r]+)',
            r'ACTION\s+NO[.:]\s*([^\n\r]+)',
        ],
        'judicial_centre': [
            r'JUDICIAL\s+CENTRE[:\s]*([^\n\r]+)',
            r'JUDICIAL\s+DISTRICT[:\s]*([^\n\r]+)',
            r'AT\s+(EDMONTON|CALGARY|LETHBRIDGE|MEDICINE\s+HAT|RED\s+DEER|ST\.\s*PAUL|WETASKIWIN)',
        ],
        'plaintiff': [
            r'PLAINTIFF[:\s]*([^\n\r]+)',
            r'PLAINTIFFS?[:\s]*([^\n\r]+)',
        ],
        'defendant': [
            r'DEFENDANT[:\s]*([^\n\r]+)',
            r'DEFENDANTS?[:\s]*([^\n\r]+)',
        ],
        'prosecution': [
            r'PROSECUTION[:\s]*([^\n\r]+)',
            r'CROWN[:\s]*([^\n\r]+)',
            r'R\.\s*V\.\s*([^\n\r]+)',  # R. v. Accused
            r'REX\s+V\.\s*([^\n\r]+)',
            r'THE\s+KING\s+V\.\s*([^\n\r]+)',
        ],
        'accused': [
            r'ACCUSED[:\s]*([^\n\r]+)',
        ],
        'document_type': [
            r'DOCUMENT[:\s]*([^\n\r]+)',
            r'TYPE\s+OF\s+DOCUMENT[:\s]*([^\n\r]+)',
        ],
        'address_for_service': [
            r'ADDRESS\s+FOR\s+SERVICE[:\s]*([^\n\r]+(?:\n[^\n\r]+){0,3})',
            r'CONTACT\s+INFORMATION[:\s]*([^\n\r]+(?:\n[^\n\r]+){0,3})',
        ],
        'date_filed': [
            r'FILED\s+(?:ON\s+)?([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'DATE\s+(?:FILED|STAMPED)[:\s]*([^\n\r]+)',
            r'CLERK\'?S?\s+STAMP[:\s]*([^\n\r]+)',
        ],
        'clerk_stamp': [
            r'CLERK\'?S?\s+STAMP[:\s]*([^\n\r]+)',
        ],
    }
    
    @classmethod
    def parse(cls, text: str, filepath: str = "") -> CKBDocument:
        """Parse document text and extract CKB header"""
        header = CKBHeader(raw_header=text[:2000])  # Store first 2000 chars as raw header
        
        # Search original text (case preserved) with IGNORECASE flag
        for field_name, patterns in cls.FIELD_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    value = match.group(1).strip()
                    # Clean up common artifacts
                    value = re.sub(r'^\s*[:\-]\s*', '', value)
                    value = re.sub(r'\s+', ' ', value)
                    setattr(header, field_name, value)
                    break
        
        # Normalize judicial centre to proper case
        if header.judicial_centre:
            centre_upper = header.judicial_centre.upper()
            for known_centre in JUDICIAL_CENTRES:
                if known_centre.upper() in centre_upper:
                    header.judicial_centre = known_centre
                    break
        
        # Special handling for judicial centre from "IN THE COURT OF KING'S BENCH"
        if not header.judicial_centre:
            centre_match = re.search(
                r'IN\s+THE\s+COURT\s+OF\s+KING\'?S?\s+BENCH\s+(?:OF\s+ALBERTA\s+)?(?:JUDICIAL\s+CENTRE\s+)?(?:OF\s+)?(EDMONTON|CALGARY|LETHBRIDGE|MEDICINE\s+HAT|RED\s+DEER|ST\.\s*PAUL|WETASKIWIN)',
                text, re.IGNORECASE
            )
            if centre_match:
                header.judicial_centre = centre_match.group(1).strip().title()
        
        # Handle accused: only set from DEFENDANT pattern if this is a criminal matter
        if header.prosecution and header.defendant and not header.accused:
            header.accused = header.defendant
            header.defendant = ""  # Clear defendant since it's actually accused in criminal context
        
        # Re-parse computed fields now that all raw fields are set
        header._parse_computed_fields()
        
        # Try to extract from filename if court file number missing
        if not header.court_file_number and filepath:
            filename = Path(filepath).name
            file_match = re.search(r'(\d{2,4}[-:]\d+)', filename)
            if file_match:
                header.court_file_number = file_match.group(1)
        
        filename = Path(filepath).name if filepath else ""
        
        return CKBDocument(
            header=header,
            content=text,
            filepath=filepath,
            filename=filename
        )


class CKBDocumentSorter:
    """Main class for sorting and formatting CKB documents"""
    
    def __init__(self, input_dir: str, output_dir: str):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.documents: List[CKBDocument] = []
        self.parser = CKBDocumentParser()
    
    def load_documents(self, extensions: List[str] = None) -> int:
        """Load and parse all documents from input directory"""
        if extensions is None:
            extensions = ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.md']
        
        count = 0
        for ext in extensions:
            for filepath in self.input_dir.rglob(f'*{ext}'):
                try:
                    text = self._extract_text(filepath)
                    if text:
                        doc = self.parser.parse(text, str(filepath))
                        self.documents.append(doc)
                        count += 1
                except Exception as e:
                    print(f"Error processing {filepath}: {e}")
        
        print(f"Loaded {count} documents")
        return count
    
    def _extract_text(self, filepath: Path) -> str:
        """Extract text from various file formats"""
        ext = filepath.suffix.lower()
        
        if ext == '.txt' or ext == '.md':
            return filepath.read_text(encoding='utf-8', errors='ignore')
        
        elif ext == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = ""
                    for page in pdf.pages[:5]:  # Only first 5 pages for header
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    return text
            except ImportError:
                print(f"pdfplumber not installed, skipping {filepath}")
                return ""
        
        elif ext in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs[:50]])  # First 50 paragraphs
                return text
            except ImportError:
                print(f"python-docx not installed, skipping {filepath}")
                return ""
        
        elif ext == '.rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                return rtf_to_text(filepath.read_text(encoding='utf-8', errors='ignore'))
            except ImportError:
                print(f"striprtf not installed, skipping {filepath}")
                return ""
        
        return ""
    
    def sort_documents(self) -> List[CKBDocument]:
        """Sort documents per CKB rules"""
        self.documents.sort(key=lambda d: d.sort_key())
        return self.documents
    
    def export_sorted(self, format: str = 'md') -> str:
        """Export sorted documents in specified format"""
        if format == 'json':
            return self._export_json()
        elif format == 'html':
            return self._export_html()
        else:
            return self._export_markdown()
    
    def _export_markdown(self) -> str:
        """Export as Markdown with CKB-compliant formatting"""
        lines = [
            "# Alberta Court of King's Bench - Sorted Document Index",
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*",
            f"*Total Documents: {len(self.documents)}*",
            "",
            "---",
            "",
        ]
        
        current_centre = None
        current_jurisdiction = None
        
        for i, doc in enumerate(self.documents, 1):
            h = doc.header
            
            # Section headers for judicial centre
            if h.judicial_centre != current_centre:
                current_centre = h.judicial_centre
                lines.append(f"\n## Judicial Centre: {current_centre or 'Unknown'}\n")
            
            # Sub-section for jurisdiction
            if h.jurisdiction != current_jurisdiction:
                current_jurisdiction = h.jurisdiction
                lines.append(f"\n### Jurisdiction: {current_jurisdiction.value.title()}\n")
            
            # Document entry with full CKB header
            lines.append(f"#### {i}. {doc.filename}")
            lines.append("")
            lines.append("| Field | Value |")
            lines.append("|-------|-------|")
            lines.append(f"| **COURT FILE NUMBER** | {h.court_file_number or 'N/A'} |")
            lines.append(f"| **JUDICIAL CENTRE** | {h.judicial_centre or 'N/A'} |")
            
            # Party fields based on jurisdiction
            if h.jurisdiction == Jurisdiction.CRIMINAL:
                lines.append(f"| **PROSECUTION** | {h.prosecution or 'N/A'} |")
                lines.append(f"| **ACCUSED** | {h.accused or h.defendant or 'N/A'} |")
            else:
                lines.append(f"| **PLAINTIFF** | {h.plaintiff or 'N/A'} |")
                lines.append(f"| **DEFENDANT** | {h.defendant or 'N/A'} |")
            
            lines.append(f"| **DOCUMENT** | {h.document_type or 'N/A'} |")
            lines.append(f"| **ADDRESS FOR SERVICE** | {h.address_for_service or 'N/A'} |")
            lines.append(f"| **DATE FILED** | {h.date_filed or 'N/A'} |")
            lines.append(f"| **CLERK'S STAMP** | {h.clerk_stamp or 'N/A'} |")
            lines.append("")
            
            # Metadata
            lines.append(f"**Metadata:** Date ID: `{h.date_id}` | Case ID: `{h.case_id}` | Year: `{h.year}` | Seq: `{h.sequence}`")
            lines.append(f"**Parsed Jurisdiction:** {h.jurisdiction.value} | **Party Type:** {h.party_type.value} | **Doc Type:** {h.doc_type_enum.value}")
            lines.append(f"**Source:** `{doc.filepath}`")
            lines.append("")
            lines.append("---")
            lines.append("")
        
        return "\n".join(lines)
    
    def _export_json(self) -> str:
        """Export as JSON"""
        data = {
            "generated": datetime.now().isoformat(),
            "total_documents": len(self.documents),
            "documents": []
        }
        
        for doc in self.documents:
            h = doc.header
            data["documents"].append({
                "filename": doc.filename,
                "filepath": doc.filepath,
                "header": {
                    "court_file_number": h.court_file_number,
                    "judicial_centre": h.judicial_centre,
                    "plaintiff": h.plaintiff,
                    "defendant": h.defendant,
                    "prosecution": h.prosecution,
                    "accused": h.accused,
                    "document_type": h.document_type,
                    "address_for_service": h.address_for_service,
                    "date_filed": h.date_filed,
                    "clerk_stamp": h.clerk_stamp,
                    "jurisdiction": h.jurisdiction.value,
                    "party_type": h.party_type.value,
                    "document_type_enum": h.doc_type_enum.value,
                    "date_id": h.date_id,
                    "case_id": h.case_id,
                    "year": h.year,
                    "sequence": h.sequence,
                },
                "sort_key": doc.sort_key(),
            })
        
        return json.dumps(data, indent=2)
    
    def _export_html(self) -> str:
        """Export as HTML table"""
        html = [
            "<!DOCTYPE html>",
            "<html><head>",
            "<title>CKB Document Index</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 20px; }",
            "table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #1a3c5e; color: white; }",
            "tr:nth-child(even) { background-color: #f2f2f2; }",
            ".centre-header { background-color: #2c5f8a; color: white; }",
            ".jurisdiction-header { background-color: #3a7bc8; color: white; }",
            ".metadata { font-size: 0.85em; color: #666; }",
            "</style>",
            "</head><body>",
            f"<h1>Alberta Court of King's Bench - Sorted Document Index</h1>",
            f"<p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total: {len(self.documents)}</p>",
        ]
        
        current_centre = None
        current_jurisdiction = None
        
        for i, doc in enumerate(self.documents, 1):
            h = doc.header
            
            if h.judicial_centre != current_centre:
                if current_centre is not None:
                    html.append("</table>")
                current_centre = h.judicial_centre
                html.append(f'<h2 class="centre-header">Judicial Centre: {current_centre or "Unknown"}</h2>')
                html.append("<table>")
                html.append("<tr><th>#</th><th>File</th><th>Court File No</th><th>Parties</th><th>Document</th><th>Date Filed</th><th>Jurisdiction</th></tr>")
            
            if h.jurisdiction != current_jurisdiction:
                current_jurisdiction = h.jurisdiction
                html.append(f'<tr class="jurisdiction-header"><td colspan="7">Jurisdiction: {current_jurisdiction.value.title()}</td></tr>')
            
            # Party display
            if h.jurisdiction == Jurisdiction.CRIMINAL:
                parties = f"Prosecution: {h.prosecution or 'N/A'}<br>Accused: {h.accused or h.defendant or 'N/A'}"
            else:
                parties = f"Plaintiff: {h.plaintiff or 'N/A'}<br>Defendant: {h.defendant or 'N/A'}"
            
            html.append(f"""
            <tr>
                <td>{i}</td>
                <td>{doc.filename}</td>
                <td>{h.court_file_number or 'N/A'}</td>
                <td>{parties}</td>
                <td>{h.document_type or 'N/A'}</td>
                <td>{h.date_filed or 'N/A'}</td>
                <td>{h.jurisdiction.value}</td>
            </tr>
            """)
        
        if current_centre is not None:
            html.append("</table>")
        
        html.append("</body></html>")
        return "\n".join(html)
    
    def copy_sorted_files(self, output_subdir: str = "sorted") -> int:
        """Copy files to output directory in sorted order with numbered prefixes"""
        sorted_dir = self.output_dir / output_subdir
        sorted_dir.mkdir(parents=True, exist_ok=True)
        
        count = 0
        for i, doc in enumerate(self.documents, 1):
            if doc.filepath and Path(doc.filepath).exists():
                src = Path(doc.filepath)
                h = doc.header
                # Create sorted filename: 001_Calgary_2303-12345_StatementOfClaim.pdf
                centre = (h.judicial_centre or "Unknown").replace(" ", "_")
                file_num = (h.court_file_number or "Unknown").replace("/", "-").replace(":", "-")
                doc_type = (h.document_type or "Document").replace(" ", "")[:30]
                ext = src.suffix
                new_name = f"{i:03d}_{centre}_{file_num}_{doc_type}{ext}"
                dst = sorted_dir / new_name
                shutil.copy2(src, dst)
                count += 1
        
        print(f"Copied {count} files to {sorted_dir}")
        return count


def main():
    parser = argparse.ArgumentParser(
        description="Alberta Court of King's Bench Document Sorter & Formatter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Sort documents in ./input to ./output as markdown
  python ckb_document_sorter.py ./input ./output --format md

  # Export as JSON for programmatic use
  python ckb_document_sorter.py ./docs ./sorted --format json

  # Copy files in sorted order with numbered prefixes
  python ckb_document_sorter.py ./input ./output --copy

  # Full pipeline: parse, sort, export all formats, copy files
  python ckb_document_sorter.py ./legal_docs ./ckb_output --format all --copy
        """
    )
    
    parser.add_argument('input_dir', help='Input directory containing legal documents')
    parser.add_argument('output_dir', help='Output directory for sorted results')
    parser.add_argument('--format', choices=['md', 'json', 'html', 'all'], default='md',
                        help='Output format (default: md)')
    parser.add_argument('--copy', action='store_true',
                        help='Copy files to output/sorted in sorted order')
    parser.add_argument('--extensions', nargs='+', 
                        default=['.pdf', '.docx', '.doc', '.txt', '.rtf', '.md'],
                        help='File extensions to process')
    
    args = parser.parse_args()
    
    # Validate input directory
    if not Path(args.input_dir).exists():
        print(f"Error: Input directory '{args.input_dir}' does not exist")
        return 1
    
    # Create output directory
    Path(args.output_dir).mkdir(parents=True, exist_ok=True)
    
    # Initialize sorter
    sorter = CKBDocumentSorter(args.input_dir, args.output_dir)
    
    # Load and parse documents
    print(f"Loading documents from {args.input_dir}...")
    sorter.load_documents(args.extensions)
    
    if not sorter.documents:
        print("No documents found or parsed successfully")
        return 1
    
    # Sort documents
    print("Sorting documents per CKB rules...")
    sorter.sort_documents()
    
    # Export in requested format(s)
    formats = ['md', 'json', 'html'] if args.format == 'all' else [args.format]
    
    for fmt in formats:
        output_file = Path(args.output_dir) / f"ckb_sorted_index.{fmt}"
        content = sorter.export_sorted(fmt)
        output_file.write_text(content, encoding='utf-8')
        print(f"Exported {fmt.upper()} to {output_file}")
    
    # Copy files if requested
    if args.copy:
        sorter.copy_sorted_files()
    
    # Print summary
    print("\n" + "="*60)
    print("SORTING COMPLETE - CKB COMPLIANT ORDER")
    print("="*60)
    print(f"Judicial Centres: {', '.join(JUDICIAL_CENTRES)}")
    print(f"Document Types: {len(DocumentType)} precedence levels")
    print(f"Jurisdictions: {', '.join([j.value for j in Jurisdiction if j != Jurisdiction.UNKNOWN])}")
    print(f"\nSort Order Applied:")
    print("  1. Judicial Centre (alphabetical)")
    print("  2. Date ID (year-sequence)")
    print("  3. Case ID (sequence)")
    print("  4. Party Type (Prosecution/Plaintiff first)")
    print("  5. Document Type (CKB precedence)")
    print("  6. Jurisdiction (Criminal > Civil > Family > Surrogate > Bankruptcy)")
    
    return 0


if __name__ == "__main__":
    exit(main())
